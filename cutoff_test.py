import xarray as xr
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from scipy.stats import gamma
from docx import Document

doc = Document()

ds = xr.open_dataset('Ts300_warm_precip.nc', decode_times=False)
prec_mp = ds['prec_mp']

flattened = prec_mp.values.flatten()
flattened = flattened[flattened > 0]
log_precip = np.log10(flattened)

max_precip = np.max(flattened)
print(f"The maximum precipitation value across all gridpoints and times: {max_precip}")

bins = np.logspace(np.min(log_precip), np.max(log_precip), 50)
plt.hist(flattened, bins=bins, density=True, alpha=0.6, color='g')

shape, loc, scale = gamma.fit(flattened)
support = np.logspace(np.log10(flattened.min()), np.log10(flattened.max()), 100)
pdf_gamma = gamma.pdf(support, shape, loc, scale)
plt.plot(support, pdf_gamma, "r-", label=f'gamma pdf')

plt.gca().set_xscale("log")
plt.gca().set_yscale("log")
plt.title('Log-Log Histogram with Gamma Fit')
plt.xlabel('Log Surface Total Precipitation Rate')
plt.ylabel('Density')
plt.legend()
plt.savefig('plot1.png')
doc.add_paragraph("This plot shows the distribution of surface total precipitation rate (log scale) with a gamma distribution fit. It's a log-log plot with the x-axis representing the log of surface total precipitation rate, and the y-axis representing the density. This histogram and the fitted gamma distribution can give us insights about the distribution of precipitation rates, showing where the majority of the values lie and how much variability there is in the data.")
doc.add_picture('plot1.png')

doc.save('document.docx')
