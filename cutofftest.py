import numpy as np
import xarray as xr
import matplotlib.pyplot as plt
import scipy.ndimage as ndi
import scipy.stats as stats
from docx import Document
from scipy.stats import gamma

def compute_cutoff_moment_ratio(cluster_array):
    mean_value = cluster_array.mean().values
    second_moment = (cluster_array ** 2).mean().values
    cutoff = second_moment / mean_value
    return cutoff

def compute_cutoff_linear_regression(cluster_array):
    hist, bins = np.histogram(cluster_array.values, bins=100, density=True)
    log_hist = np.log(hist)
    valid = log_hist != -np.inf
    valid_bins = bins[:-1][valid]
    valid_log_hist = log_hist[valid]
    slope, intercept, _, _, _ = stats.linregress(valid_bins[-20:], valid_log_hist[-20:])
    cutoff = -1 / slope
    return cutoff

files = ['file1.nc', 'file2.nc', 'file3.nc']  

cutoffs_moment_ratio = []
cutoffs_linear_regression = []

for file in files:

    doc = Document()
    ds = xr.open_dataset(file, decode_times=False)
    prec_mp = ds['prec_mp']

    flattened = prec_mp.values.flatten()
    flattened = flattened[flattened > 0]
    log_precip = np.log10(flattened)

    cutoff_moment_ratio = compute_cutoff_moment_ratio(prec_mp)
    cutoff_linear_regression = compute_cutoff_linear_regression(prec_mp)

    cutoffs_moment_ratio.append(cutoff_moment_ratio)
    cutoffs_linear_regression.append(cutoff_linear_regression)

    print(f"File: {file}")
    print(f"Cutoff (Moment Ratio Method): {cutoff_moment_ratio}")
    print(f"Cutoff (Linear Regression Method): {cutoff_linear_regression}")

    max_precip = np.max(flattened)
    print(f"The maximum precipitation value across all gridpoints and times: {max_precip}")

    plt.figure()

    # Histogram
    bins = np.logspace(np.min(log_precip), np.max(log_precip), 50)
    plt.hist(flattened, bins=bins, density=True, alpha=0.6, color='g')

    shape, loc, scale = gamma.fit(flattened)
    support = np.logspace(np.log10(flattened.min()), np.log10(flattened.max()), 100)
    pdf_gamma = gamma.pdf(support, shape, loc, scale)
    plt.plot(support, pdf_gamma, "r-", label=f'gamma pdf')

    plt.gca().set_xscale("log")
    plt.gca().set_yscale("log")
    plt.title('Log-Log Histogram with Gamma Fit')
    plt.xlabel('Log Surface Total Precipitation Rate')
    plt.ylabel('Density')
    plt.legend()
    plt.savefig(f'{file}_plot.png')

    doc.add_paragraph(f"This plot shows the distribution of surface total precipitation rate (log scale) for the file {file} with a gamma distribution fit.")
    doc.add_paragraph(f"Cutoff (Moment Ratio Method): {cutoff_moment_ratio}")
    doc.add_paragraph(f"Cutoff (Linear Regression Method): {cutoff_linear_regression}")
    doc.add_paragraph(f"The maximum precipitation value across all gridpoints and times: {max_precip}")

    doc.add_picture(f'{file}_plot.png')

    doc.save(f'{file}_document.docx')

plt.figure()
plt.plot(cutoffs_moment_ratio, label="Moment Ratio Method")
plt.plot(cutoffs_linear_regression, label="Linear Regression Method")
plt.xlabel('File index')
plt.ylabel('Cutoff value')
plt.legend()
plt.show()

