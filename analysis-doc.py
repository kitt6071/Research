#!/usr/bin/env python
# coding: utf-8

# In[ ]:


import xarray as xr
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import hvplot.xarray
from hvplot import hvPlot
from scipy.stats import gamma
from docx import Document
doc = Document()

ds = xr.open_dataset('Ts300_warm_dx0.125_npz151_3d_inst.nc', decode_times=False)
time_range = ds.time.values
grid_xt_range = ds.grid_xt.values
grid_yt_range = ds.grid_yt.values

prec_mp = ds['prec_mp']

flattened = prec_mp.values.flatten()
flattened = flattened[flattened > 0]
log_precip = np.log10(flattened)

# Histogram
bins = np.logspace(np.min(log_precip), np.max(log_precip), 50)
plt.hist(flattened, bins=bins, density=True, alpha=0.6, color='g')


# Fit a gamma distribution
shape, loc, scale = gamma.fit(flattened)
support = np.logspace(np.log10(flattened.min()), np.log10(flattened.max()), 100)
pdf_gamma = gamma.pdf(support, shape, loc, scale)
plt.plot(support, pdf_gamma, "r-", label=f'gamma pdf')

# axes to log scale
plt.gca().set_xscale("log")
plt.gca().set_yscale("log")
plt.title('Log-Log Histogram with Gamma Fit')
plt.xlabel('Log Surface Total Precipitation Rate')
plt.ylabel('Density')
plt.legend()
plt.show()
flattened = prec_mp.values.flatten()
plt.savefig('plot1.png')
doc.add_paragraph("This plot shows the distribution of surface total precipitation rate (log scale) with a gamma distribution fit. It's a log-log plot with the x-axis representing the log of surface total precipitation rate, and the y-axis representing the density. This histogram and the fitted gamma distribution can give us insights about the distribution of precipitation rates, showing where the majority of the values lie and how much variability there is in the data.")
doc.add_picture('plot1.png')

plt.hist(flattened, bins=100)
plt.title('Histogram of Surface Total Precipitation Rate')
plt.xlabel('Surface Total Precipitation Rate (mm/day)')
plt.ylabel('Frequency')
plt.show()
plt.savefig('plot2.png')
doc.add_paragraph('This plot shows a histogram of the surface total precipitation rate. The x-axis represents the surface total precipitation rate in mm/day, and the y-axis represents the frequency. This plot provides a straightforward way to visualize how often different ranges of precipitation rates occur, giving us an understanding of common, rare, and extreme precipitation rates.')
doc.add_picture('plot2.png')



# In[ ]:


print(ds.variables)
precip_rate = ds['prec_mp']
# ravel flattens to 1d array
precip_rate_1D = precip_rate.values.ravel()
# remove invalid data
precip_rate_1D = precip_rate_1D[~np.isnan(precip_rate_1D)]
plt.hist(precip_rate_1D, bins=60) 
plt.show()
plt.savefig('plot3.png')
doc.add_paragraph("This histogram displays the distribution of the precipitation rates after they've been transformed into a one-dimensional array. By disregarding the spatial structure of the data, this histogram shows the overall distribution of precipitation rates in the dataset.")

doc.add_picture('plot3.png')



# In[ ]:


airr = ds.prec_mp
air1d = airr.isel(grid_xt=10, grid_yt=10)

air1d.plot()
plt.show()
plt.savefig('plot4.png')
doc.add_paragraph("This line graph shows how the precipitation rate at a certain location (grid_xt=10, grid_yt=10) changes over time. This can provide insights about any temporal patterns or trends in the precipitation rate at this location, such as periodic cycles or long-term increases or decreases.")

doc.add_picture('plot4.png')

air1d[:200].plot.line("b-^")
air1d.plot()
plt.show()


# In[ ]:


prec_at_loc = ds['prec_mp'].isel(grid_xt=10, grid_yt=10)

# time-series graph of the precipitation rate at location
prec_at_loc.plot()
plt.show()
plt.savefig('plot6.png')
doc.add_paragraph("This plot shows the precipitation rate at a certain location (grid_xt=10, grid_yt=10) over time. It's a time-series plot which provides insight into how precipitation changes over time at this specific location, revealing trends, patterns or outliers.")

doc.add_picture('plot6.png')
prec_at_time = ds['prec_mp'].isel(time=0)


# In[ ]:


# histogram of the precipitation rates
prec_at_time.plot.hist(bins=100)
plt.show()

prec_at_time.plot()
plt.show()
prec_at_time.plot.hist(bins=30, range=[0, 20])
plt.show()
plt.savefig('plot7.png')
doc.add_paragraph("This histogram shows the distribution of precipitation rates at a specific time (time=0), specifically focusing on the range 0-20 mm/day. This plot allows us to see the frequency of different precipitation rates during this specific time period.")

doc.add_picture('plot7.png')


# In[ ]:


prec_at_loc = ds['prec_mp'].isel(grid_xt=10, grid_yt=10)
# 3D graph!pip install hvplot
prec_at_loc.hvplot()
# 2D plot
# surface total precipitation rate every 10 step
prec_at_times = ds['prec_mp'].isel(time=slice(0,60,10)) 
prec_at_times.plot(col='time', col_wrap=3)
plt.show()
plt.savefig('plot8.png')
doc.add_paragraph("This set of plots shows the spatial distribution of precipitation rates at different times, specifically every 10th time step. Each plot provides a snapshot of the precipitation rates at one time, allowing us to see how the spatial pattern of precipitation changes over time.")

doc.add_picture('plot8.png')


# In[ ]:


# average over spatial dimensions
average_prec = ds['prec_mp'].mean(dim=['grid_xt', 'grid_yt'])

# line plot
average_prec.plot()
plt.title('Average Precipitation Rate Over Time')
plt.xlabel('Time')
plt.ylabel('Average Precipitation Rate (mm/day)')
plt.show()
plt.savefig('plot9.png')
doc.add_paragraph("This line graph shows the average precipitation rate over time. It provides a summary of how the overall level of precipitation changes over time, smoothing out the local and regional variability.")

doc.add_picture('plot9.png')








# In[ ]:


# group by lat bins and calc mean and std dev
prec_mp_grp = ds['prec_mp'].mean(["time", "grid_xt"]).groupby_bins("grid_yt", [0, 23.5, 66.5, 90])
prec_mp_mean = prec_mp_grp.mean()
prec_mp_std = prec_mp_grp.std()
# mean and the mean +/- std dev
prec_mp_mean.plot.step()
(prec_mp_mean + prec_mp_std).plot.step(ls=":")
(prec_mp_mean - prec_mp_std).plot.step(ls=":")

# plot limits and title
plt.ylim(0, 90)
plt.title("Zonal mean precipitation rate")
plt.show()
plt.savefig('plot18.png')
doc.add_paragraph("This step plot shows the zonal mean precipitation rate and its standard deviation. The plot provides insight into how precipitation varies by latitude, indicating which latitudinal zones have higher or lower average precipitation.")

doc.add_picture('plot18.png')



# In[ ]:


# plot against latitude
ds['prec_mp'].isel(time=10, grid_xt=[10, 11]).plot(y="grid_yt", hue="grid_xt")
plt.show()
plt.savefig('plot19.png')
doc.add_paragraph("This plot shows the precipitation rates at time=10 for two different longitude points (grid_xt=10, 11), plotted against latitude. It provides a way to compare the latitudinal profile of precipitation at these two points.")

doc.add_picture('plot19.png')



# In[ ]:


# plot against time
ds['prec_mp'].isel(grid_xt=10, grid_yt=[19, 21, 22]).plot.line(x="time")
plt.show()
plt.savefig('plot20.png')
doc.add_paragraph("This line plot shows the precipitation rates at a specific longitude point (grid_xt=10) for three different latitude points (grid_yt=19, 21, 22), plotted over time. It provides a means to compare the temporal pattern of precipitation at these three points.")

doc.add_picture('plot20.png')


# In[ ]:


# precipitation rate at location over the first 10 steps
ds['prec_mp'].isel(grid_xt=10, grid_yt=10, time=slice(0,10)).plot.line(drawstyle='steps')
plt.show()
plt.savefig('plot14.png')
doc.add_paragraph("This step plot shows the precipitation rate at a certain location (grid_xt=10, grid_yt=10) for the first 10 steps. It provides a clear view of the changes in precipitation rate during these initial time steps.")

doc.add_picture('plot14.png')


# In[ ]:


# line plot with log y
ds['prec_mp'].isel(grid_xt=10, grid_yt=10, time=slice(0,60)).plot.line(yscale='log')
plt.show()
plt.savefig('plot15.png')
doc.add_paragraph("This line plot shows the precipitation rate at a certain location (grid_xt=10, grid_yt=10) over time, with the y-axis on a log scale. This can emphasize relative changes and is particularly useful when the data spans several orders of magnitude.")

doc.add_picture('plot15.png')



# In[ ]:


# surface plot precipitation rate
ds['prec_mp'].isel(time=0).plot.surface()
plt.show()
plt.savefig('plot16.png')
doc.add_paragraph("This surface plot depicts the precipitation rate at time=0 in a 3D perspective. This representation helps understand patterns, peaks, and valleys of precipitation rate in a spatial context.")

doc.add_picture('plot16.png')


# In[ ]:


# facet plot of precipitation rate
ds['prec_mp'].isel(time=slice(0,60,10)).plot(col='time', col_wrap=3)
plt.show()
plt.savefig('plot17.png')
doc.add_paragraph("These facet plots represent the precipitation rate at different time steps, specifically every 10th time step. Each plot is a snapshot of the precipitation rate at a specific time, enabling comparisons of the spatial pattern of precipitation at different times.")

doc.add_picture('plot17.png')


# In[ ]:


prec_at_time = ds['prec_mp'].isel(time=0)

# contours of precipitation rate
prec_at_time.plot.contour(levels=10)
plt.show()
plt.savefig('plot10.png')
doc.add_paragraph("This contour plot shows the spatial distribution of precipitation rates at a certain time (time=0). The contours allow us to see regions of similar precipitation rates, providing a summary of the spatial pattern of precipitation.")

doc.add_picture('plot10.png')



# In[ ]:


prec_at_time = ds['prec_mp'].isel(time=0)

# pcolormesh plot
prec_at_time.plot.pcolormesh()
plt.show()
plt.savefig('plot11.png')
doc.add_paragraph("This is a pseudocolor plot which displays the precipitation rate at a certain time (time=0) over a 2D grid. The color intensity corresponds to the precipitation rate. It's a great tool to visualize spatial patterns and areas with higher or lower precipitation rates.")

doc.add_picture('plot11.png')


# In[ ]:


ds['prec_mp'].isel(grid_xt=10, grid_yt=10, time=slice(0,10)).plot.line(x='time', hue='time')
plt.show()
plt.savefig('plot12.png')
doc.add_paragraph("This line plot shows the precipitation rate at a certain location (grid_xt=10, grid_yt=10) for the first 10 steps. It provides a quick snapshot of how precipitation changes over the early period of time.")

doc.add_picture('plot12.png')


# In[ ]:


# precipitation rates with time on y 
ds['prec_mp'].isel(grid_xt=10, grid_yt=10, time=slice(time_range[0], time_range[-1])).plot.line(y='time')
plt.show()
plt.savefig('plot13.png')
doc.add_paragraph("This line plot depicts precipitation rates at a particular location (grid_xt=10, grid_yt=10) with time represented on the y-axis. It offers a different perspective on how precipitation rates change over time.")

doc.add_picture('plot13.png')

